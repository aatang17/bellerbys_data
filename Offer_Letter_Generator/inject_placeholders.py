"""
One-time script: read the current template and write a new copy with {{Placeholder}} inserted
so the generator can replace them. Run from project root:
  python inject_placeholders.py
Output: templates/offer_letter_template_with_placeholders.docx
Then either rename it to offer_letter_template.docx or set OFFER_TEMPLATE_PATH to the new file.
"""
import os
from pathlib import Path
from docx import Document

BASE = Path(__file__).parent
TEMPLATE = BASE / "templates" / "offer_letter_template.docx"
OUT = BASE / "templates" / "offer_letter_template_with_placeholders.docx"

REPLACEMENTS = [
    ("2601270", "{{Student_ID}}"),
    ("Kanta Koizumi  小泉環太", "{{Name_En}}  {{Name_Zh}}"),
    ("Kanta Koizumi", "{{Name_En}}"),
    ("小泉環太", "{{Name_Zh}}"),
    ("29 November 2007", "{{DOB}}"),
    ("2007 / 11 / 29", "{{DOB_Short}}"),
    ("kantakoizumi001@gmail.com", "{{Email}}"),
    ("Media Foundation Program  传媒预科课程", "{{Program}}"),
    ("Beijing Normal-Hong Kong Baptist University (BNBU) 北师香港浸会大学", "{{Campus}}"),
    ("14 September 2026", "{{Commencement}}"),
    ("30 June 2027", "{{Expected_Graduation}}"),
    ("CNY 20,000   （贰万元整）（Deductible from tuition fee）", "{{Enrollment_Deposit}}"),
    ("1 May 2026", "{{Deposit_Deadline}}"),
    ("CNY 158,000  （壹拾伍万捌仟元整）", "{{Tuition_Fee}}"),
    ("1 August 2026", "{{Tuition_Deadline}}"),
    ("¥158,000", "{{Tuition_Amount}}"),
    ("¥10,000", "{{Dormitory_Amount}}"),
    ("¥20,000", "{{Scholarship_Amount}}"),
    ("¥148,000", "{{Total_Amount}}"),
    ("9 March 2026", "{{Issue_Date}}"),
    ("23 March 2026", "{{Payment_Deadline}}"),
    ("2026年3月9日", "{{Issue_Date_ZH}}"),
    ("2026年3月23日", "{{Payment_Deadline_ZH}}"),
]


def replace_in_paragraph(paragraph):
    for old, new in REPLACEMENTS:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)


def replace_in_cell(cell):
    for p in cell.paragraphs:
        replace_in_paragraph(p)


def main():
    if not TEMPLATE.is_file():
        print(f"Template not found: {TEMPLATE}")
        return
    doc = Document(TEMPLATE)
    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Written: {OUT}")
    print("Rename to offer_letter_template.docx or set env OFFER_TEMPLATE_PATH to this file.")


if __name__ == "__main__":
    main()
