"""
Generate offer letter DOCX from template by replacing {{Placeholder}} and optional scholarship block.
"""
import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document

# Month names for DOB formatting
_MONTHS = ["", "January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]


def _current_issue_dates() -> tuple[str, str]:
    """Return (Issue_Date, Issue_Date_ZH) using today's date for page 1 and page 2."""
    now = datetime.now()
    en = f"{now.day} {_MONTHS[now.month]} {now.year}"
    zh = f"{now.year}年{now.month}月{now.day}日"
    return en, zh


def _format_dob(dob: str) -> tuple[str, str]:
    """Convert YYYY-MM-DD or '29 November 2007' to (long form, short form)."""
    if not dob or not dob.strip():
        return "", ""
    dob = dob.strip()
    # Already long form (e.g. 29 November 2007)
    if " " in dob and not re.match(r"^\d{4}-\d{2}-\d{2}$", dob):
        short = dob.replace(" ", " / ")
        return dob, short
    # YYYY-MM-DD from date input
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", dob)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), m.group(3)
        if 1 <= mo <= 12:
            long_form = f"{int(d)} {_MONTHS[mo]} {y}"
            short_form = f"{y} / {mo} / {int(d)}"
            return long_form, short_form
    return dob, dob.replace(" ", " / ")


BASE = os.path.dirname(os.path.abspath(__file__))
def _default_template_path():
    with_ph = os.path.join(BASE, "templates", "offer_letter_template_with_placeholders.docx")
    if os.path.isfile(with_ph):
        return with_ph
    return os.path.join(BASE, "templates", "offer_letter_template.docx")

TEMPLATE_PATH = os.environ.get("OFFER_TEMPLATE_PATH") or _default_template_path()
GENERATED_DIR = os.environ.get(
    "OFFER_GENERATED_DIR",
    os.path.join(BASE, "generated"),
)


def _replace_in_paragraph(paragraph, replacements):
    """Replace {{key}} with value in a paragraph (may span multiple runs)."""
    full = "".join(run.text for run in paragraph.runs)
    if not full.strip():
        return
    new_text = full
    for key, value in replacements.items():
        if value is None:
            value = ""
        placeholder = "{{" + key + "}}"
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(value))
    if new_text != full:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def _replace_in_cell(cell, replacements):
    for paragraph in cell.paragraphs:
        _replace_in_paragraph(paragraph, replacements)


def _remove_scholarship_block(doc):
    """Remove paragraphs that form the Scholarship Award section (EN + ZH)."""
    to_remove = []
    in_block = False
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if "Scholarship Award" in text or "奖学金授予" in text:
            in_block = True
        if in_block:
            to_remove.append(i)
            if "Accommodation" in text or "住宿" in text:
                break
    # Remove from end so indices stay valid
    for i in reversed(to_remove):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)


def _replace_scholarship_in_tables(doc, scholarship_amount):
    """In fees table, set scholarship row to amount or '—' and recalc total if needed."""
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if "Scholarship" in str(cells) or "奖学金" in str(cells):
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if "¥" in run.text or "—" in run.text:
                                run.text = run.text.replace("¥20,000", scholarship_amount or "—").replace("¥20,000", scholarship_amount or "—")
                                run.text = run.text.replace("— ¥20,000", "— " + (scholarship_amount or "—"))
                break


def get_global_vars() -> dict:
    """Load global variables from DB for use in template."""
    import db
    with db.get_db() as conn:
        rows = conn.execute("SELECT key, value FROM global_vars").fetchall()
    return {r["key"]: r["value"] for r in rows}


def generate_letter(
    student_id: str,
    student_name: str,
    name_zh: str = "",
    dob: str = "",
    email: str = "",
    program: str = "",
    scholarship_amount: str | None = None,
    scholarship_details: str | None = None,
) -> tuple[str, str]:
    """
    Fill template with student and global vars; optionally include/remove scholarship section.
    Returns (absolute path to saved DOCX, filename).
    """
    if not os.path.isfile(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)
    globals_ = get_global_vars()
    issue_date_en, issue_date_zh = _current_issue_dates()

    dob_long, dob_short = _format_dob(dob or "")
    # Build replacement dict: per-student + global; Issue_Date/Issue_Date_ZH use current date
    replacements = {
        "Student_ID": student_id,
        "Application_Number": student_id,
        "Name_En": student_name,
        "Name_Zh": name_zh or student_name,
        "DOB": dob_long,
        "DOB_Short": dob_short,
        "Email": email,
        "Program": program or globals_.get("Program", ""),
        "Campus": globals_.get("Campus", ""),
        "Commencement": globals_.get("Commencement", ""),
        "Expected_Graduation": globals_.get("Expected_Graduation", ""),
        "Enrollment_Deposit": globals_.get("Enrollment_Deposit", ""),
        "Deposit_Deadline": globals_.get("Deposit_Deadline", ""),
        "Tuition_Fee": globals_.get("Tuition_Fee", ""),
        "Tuition_Amount": globals_.get("Tuition_Amount", ""),
        "Tuition_Deadline": globals_.get("Tuition_Deadline", ""),
        "Dormitory_Amount": globals_.get("Dormitory_Amount", ""),
        "Issue_Date": issue_date_en,
        "Issue_Date_ZH": issue_date_zh,
        "Payment_Deadline": globals_.get("Payment_Deadline", ""),
        "Payment_Deadline_ZH": globals_.get("Payment_Deadline_ZH", ""),
        "Scholarship_Amount": scholarship_amount or "—",
        "Scholarship_Details": scholarship_details or "",
    }
    # Total: from globals or default (tuition + dormitory - scholarship)
    replacements["Total_Amount"] = globals_.get("Total_Amount", "¥148,000")
    # Allow global_vars to override any key
    for k, v in globals_.items():
        if k not in replacements:
            replacements[k] = v

    # Replace in all paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)

    # Scholarship: remove block if no scholarship
    if not scholarship_amount:
        _remove_scholarship_block(doc)
    else:
        _replace_scholarship_in_tables(doc, scholarship_amount)

    Path(GENERATED_DIR).mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^\w\-_.\s]", "", f"{student_id}_{student_name}".replace(" ", "_"))[:80]
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    file_name = f"{safe_name}_{ts}.docx"
    out_path = os.path.join(GENERATED_DIR, file_name)
    doc.save(out_path)
    return os.path.abspath(out_path), file_name
